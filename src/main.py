"""
legalreview - 契約書ファイル（.doc/.docx）を読み込み、法務レビューテンプレートを生成するCLIツール
"""
import argparse
import glob
import os
import shutil
import subprocess
import sys
from datetime import datetime
from pathlib import Path


def get_timestamp():
    return datetime.now().strftime("%Y%m%d_%H%M%S")


def write_file(path: str, content: str):
    Path(path).parent.mkdir(parents=True, exist_ok=True)
    with open(path, "w", encoding="utf-8") as f:
        f.write(content)


def read_docx(path: str) -> str:
    from docx import Document
    doc = Document(path)
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    # テーブル内テキストも取得
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text = cell.text.strip()
                if text:
                    paragraphs.append(text)
    return "\n\n".join(paragraphs)


def read_doc(path: str) -> str:
    if shutil.which("antiword"):
        result = subprocess.run(["antiword", path], capture_output=True, text=True)
        if result.returncode == 0 and result.stdout.strip():
            return result.stdout
    # フォールバック：バイナリからASCII可能文字を抽出（簡易）
    with open(path, "rb") as f:
        raw = f.read()
    text = raw.decode("latin-1", errors="ignore")
    import re
    lines = re.findall(r"[\x20-\x7E\u3000-\u9FFF\uFF00-\uFFEF]{5,}", text)
    return "\n".join(lines)


def extract_text(input_path: str) -> str:
    ext = Path(input_path).suffix.lower()
    if ext == ".docx":
        return read_docx(input_path)
    elif ext == ".doc":
        return read_doc(input_path)
    else:
        print(f"エラー: 対応していない拡張子です（{ext}）。.doc または .docx を指定してください。", file=sys.stderr)
        sys.exit(1)


def build_contract_md(text: str, filename: str) -> str:
    now = datetime.now().strftime("%Y-%m-%d %H:%M")
    lines = [
        f"# 契約書テキスト: {filename}",
        f"\n抽出日時: {now}",
        "\n---\n",
        "## 抽出テキスト\n",
        "```",
        text,
        "```",
    ]
    return "\n".join(lines) + "\n"


def cmd_summarize(args):
    input_path = os.path.join("input", args.file)
    if not os.path.exists(input_path):
        print(f"エラー: {input_path} が見つかりません", file=sys.stderr)
        sys.exit(1)

    text = extract_text(input_path)
    content = build_contract_md(text, args.file)

    out_path = os.path.join("mid", f"contract_{get_timestamp()}.md")
    write_file(out_path, content)
    print(f"完了: {out_path}")


def build_review_template(contract_filename: str, contract_md_path: str) -> str:
    now = datetime.now()
    date_str = now.strftime("%Y年%m月%d日 %H:%M")

    template_path = Path(__file__).parent.parent / "direction" / "review_template.md"
    template = template_path.read_text(encoding="utf-8")
    return template.format(
        contract_filename=contract_filename,
        date_str=date_str,
        contract_md_path=contract_md_path,
    )


def cmd_review(_args):
    pattern = os.path.join("mid", "contract_*.md")
    files = sorted(glob.glob(pattern))
    if not files:
        print("エラー: mid/contract_*.md が見つかりません。先に summarize を実行してください。", file=sys.stderr)
        sys.exit(1)

    latest = files[-1]

    # 元の契約書ファイル名を抽出（mid/contract_*.md の先頭行から）
    contract_filename = "（不明）"
    with open(latest, encoding="utf-8") as f:
        first_line = f.readline().strip()
    if first_line.startswith("# 契約書テキスト: "):
        contract_filename = first_line.replace("# 契約書テキスト: ", "")

    content = build_review_template(contract_filename, latest)
    out_path = os.path.join("output", f"legalreview_{get_timestamp()}.md")
    write_file(out_path, content)
    print(f"完了: {out_path}")


def main():
    parser = argparse.ArgumentParser(description="契約書法務レビューツール")
    sub = parser.add_subparsers(dest="command")

    p_sum = sub.add_parser("summarize", help="契約書ファイルからテキストを抽出する")
    p_sum.add_argument("file", help="input/ 以下のファイル名（.doc または .docx）")

    sub.add_parser("review", help="最新の契約書テキストから法務レビューテンプレートを生成する")

    args = parser.parse_args()
    if args.command == "summarize":
        cmd_summarize(args)
    elif args.command == "review":
        cmd_review(args)
    else:
        parser.print_help()
        sys.exit(1)


if __name__ == "__main__":
    main()
